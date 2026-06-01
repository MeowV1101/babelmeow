"""Build a Google-Docs-ready .docx of the BabelMeow journey (Thai)."""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(r"D:\claude\BabelMeow\docs\BabelMeow_Journey.docx")
THAI_FONT = "Tahoma"  # supports Thai + renders in Google Docs

doc = Document()

# ---- default font (incl. complex-script for Thai) ----
normal = doc.styles["Normal"]
normal.font.name = THAI_FONT
normal.font.size = Pt(11)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:cs"), THAI_FONT)   # complex script (Thai)
rfonts.set(qn("w:ascii"), THAI_FONT)
rfonts.set(qn("w:hAnsi"), THAI_FONT)


def _set_cs(run):
    r = run._element.get_or_add_rPr().get_or_add_rFonts()
    r.set(qn("w:cs"), THAI_FONT); r.set(qn("w:ascii"), THAI_FONT); r.set(qn("w:hAnsi"), THAI_FONT)


def para(text="", size=11, bold=False, color=None, align=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = color
        _set_cs(r)
    return p


def heading(text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: RGBColor(0x1F, 0x4E, 0x79), 2: RGBColor(0x2E, 0x75, 0xB6), 3: RGBColor(0x40, 0x40, 0x40)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(sizes[level]); r.font.color.rgb = colors[level]
    _set_cs(r)
    # outline level for TOC/navigation
    pPr = p._p.get_or_add_pPr(); ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), str(level-1)); pPr.append(ol)
    return p


def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(size); _set_cs(r)
    p.paragraph_format.space_after = Pt(3)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10); _set_cs(run)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(10); _set_cs(run)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ===================== CONTENT =====================
title = para("BabelMeow", size=30, bold=True, color=RGBColor(0x1F, 0x4E, 0x79),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("บันทึกการเดินทาง — เอกสารเรียนรู้สำหรับมือใหม่", size=14, color=RGBColor(0x60, 0x60, 0x60),
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("ระบบแปล Diablo IV เป็นภาษาไทย โดยไม่แตะไฟล์เกม", size=11, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

para("เอกสารนี้เล่าทุกขั้นตอนของการสร้างระบบ รวมถึงสิ่งที่ลองแล้วไม่เวิร์ก เพราะ \"ทางตัน\" คือบทเรียนที่ดีที่สุด — เป้าหมายคือให้เข้าใจ \"ทำไม\" ไม่ใช่แค่ \"ทำอะไร\"", italic=True)

heading("🎯 โจทย์เริ่มต้น")
para("\"อยากทำ mod ภาษาไทย Diablo IV\" — ฟังดูง่าย แต่พอเจาะจริงมีกำแพงเต็มไปหมด เอกสารนี้เล่าว่าชนกำแพงอะไรบ้าง และข้ามมายังไง")

heading("🗺 ภาพรวมระบบที่ได้")
para("ทำครั้งเดียว (OFFLINE): ดึงข้อความจากเกม 174k ข้อความ → แปลด้วย AI ในเครื่อง → เก็บเป็นพจนานุกรม SQLite 96k คำ", bold=True)
para("ตอนเล่น (RUNTIME): จับภาพจอ → OCR อ่านตัวอักษร → ค้นพจนานุกรม → วาด overlay ไทยทับจอ (ถ้าไม่มีในพจนานุกรม → AI แปลสด + จำไว้)", bold=True)
para("หัวใจ: แปลล่วงหน้าเก็บ cache → ตอนเล่นแค่ \"เปิดพจนานุกรม\" (เร็ว ไม่กิน VRAM) ไม่ใช่ \"คิดแปลใหม่ทุกครั้ง\" (ช้า กิน VRAM แย่งเกม)")

heading("🧱 ทำไมไม่ mod ตรงๆ — ทางตันแรก")
para("ความคิดแรก: \"ก็แก้ไฟล์ภาษาในเกมเลยสิ\" — แต่ทำไม่ได้เพราะ:")
table(["กำแพง", "เหตุผล"], [
    ["ไฟล์ถูก pack", "D4 ใช้ระบบ CASC — ทุกอย่างอยู่ในไฟล์ data.000-data.156"],
    ["online + anti-cheat", "แก้ไฟล์ client = เสี่ยงโดน ban"],
    ["launcher verify", "เปิดเกมทีก็ดาวน์โหลดไฟล์เดิมกลับ"],
    ["ฟอนต์ไม่มีไทย", "แปลได้แต่ตัวอักษรเป็น □□□ (font atlas ไม่มี glyph ไทย)"],
], widths=[1.8, 5.0])
para("บทเรียน #1: เข้าใจข้อจำกัดแพลตฟอร์มก่อน — บางวิธีตรงสุดแต่ทำไม่ได้", bold=True, color=RGBColor(0xC0,0x50,0x00))
para("ทางออก: ไม่แตะเกมเลย → ทำ overlay translator (อ่านจอ → แปล → วาดทับ) เหมือน OBS/Discord = ปลอดภัย")

heading("🔬 8 สิ่งที่ลองแล้วไม่เวิร์ก (เรียนจากความผิดพลาด)")
deadends = [
    ("1. DeepL แปลไทย", "DeepL ไม่รองรับภาษาไทย → บทเรียน: เช็ค tool รองรับภาษาเป้าหมายก่อนวางแผน"),
    ("2. แปลสดด้วย LLM ตอนเล่น", "D4 8GB + LLM 6GB + OCR = VRAM เกิน เกมกระตุก → เลยเลือกแปลล่วงหน้าแทน"),
    ("3. ใช้ Claude chat แปลทั้งหมด", "100k strings = chat 1000+ ครั้ง เปลืองมาก → ควร batch ด้วย local model"),
    ("4. Ollama รันบน CPU (ไม่รู้ตัว!)", "RX 9070 XT ใหม่เกินไป fallback ไป CPU เงียบๆ ช้า 5 เท่า → แก้ด้วย OLLAMA_VULKAN=1. บทเรียนสำคัญสุด: อย่าเชื่อว่าใช้ GPU วัดเสมอ"),
    ("5. d4-asset-extractor", "ดึงแค่ texture ไม่ดึงข้อความ → เปลี่ยนไปใช้ D4Analyzer"),
    ("6. localhost ช้า 2 วินาที", "Windows ลอง IPv6 ก่อน timeout 2 วิ → แก้ด้วย dual-stack socket. บทเรียน: ปัญหาช้าอาจอยู่ที่ network ไม่ใช่ logic"),
    ("7. live fallback ช้า 34 วินาที", "glossary prompt ยาวทำ prefill ช้าบน 4B → ใช้ light prompt เหลือ ~1 วิ (เร็วขึ้น 30 เท่า)"),
    ("8. Settle Time ไม่ใช่ตัวการ", "คิดว่า OCR ช้า แต่วัดแล้ว OCR 76ms, bridge 3ms → ตัวจริงคือ live แปลคำใหม่. บทเรียน: วัดก่อนแก้"),
]
for t_, d in deadends:
    para(t_, bold=True, color=RGBColor(0xB0,0x30,0x30), space_after=1)
    para(d, space_after=8)

heading("🪜 การเดินทางทีละ Phase")
phases = [
    ("Phase 0 — สำรวจ & ตัดสินใจ", "พบ CASC → เลือก overlay (ไม่ mod), pre-translate (ไม่ live), local AI (ฟรี/privacy). เครื่องมือ: RST + Ollama + Typhoon + Python"),
    ("Phase 1 — Pipeline แปล", "ติดตั้ง tools, โหลด Typhoon 8B/4B, เขียน glossary + post-process. ค้นพบ: glossary บนสุดของ prompt → 10/10"),
    ("Phase 2 — ดึงข้อความ", "D4Analyzer → Copy Selected → TSV → JSON. ได้ 174,817 ข้อความ (dedup เหลือ 96k ประหยัด 41%)"),
    ("Phase 3 — แปลทั้งหมด", "SQLite cache resume-able, parallel workers. ค้นพบ CPU bug → แก้เป็น GPU เร็วขึ้น 7 เท่า (80ชม→12ชม)"),
    ("Phase 4 — Bridge", "FastAPI ปลอม Ollama, 3 ชั้นค้นหา exact/normalized/fuzzy. ค้นพบ RST v5 ส่ง JSON + ##|||## separator"),
    ("Phase 5 — Dynamic text", "\"Defeat the {MONSTER}\" → template + แปลค่าต่อ. \"+12 Strength\" → \"+12 พลังกาย\". ขยาย markup ({if}/{c}/{SF})"),
    ("Phase 6 — เล่นจริง", "แก้: port, RST format, dual-stack, markup, live fallback concurrent + light prompt → เล่นได้จริง!"),
]
for t_, d in phases:
    para(t_, bold=True, color=RGBColor(0x2E,0x75,0xB6), space_after=1)
    para(d, space_after=8)

heading("💎 8 บทเรียนสำคัญ")
for b in [
    "เข้าใจข้อจำกัดแพลตฟอร์มก่อน — บางวิธีตรงสุดแต่ทำไม่ได้ (mod = ban)",
    "วัดก่อนแก้ — CPU/GPU, localhost/IPv6, OCR/bridge/live เดาผิดเสียเวลา",
    "แยกชั้นเวลา debug — ปัญหาช้าอาจอยู่ที่ network ไม่ใช่ logic",
    "pre-compute ดีกว่า compute-on-demand — แปลล่วงหน้าเก็บ cache",
    "fallback เป็นชั้นๆ — exact → fuzzy → template → live → echo (เร็วก่อน แม่นทีหลัง)",
    "memoize ทุกอย่างที่ซ้ำ — RST ส่งซ้ำตลอด จำไว้ = เร็วขึ้นมหาศาล",
    "tool ที่ใช่ ดีกว่าฝืน tool ที่ผิด — d4-asset-extractor → D4Analyzer",
    "commit บ่อยๆ — แต่ละ phase push GitHub ย้อนได้ มี CI ตรวจ",
]:
    bullet(b)

heading("⚙️ ระบบ 5 ชั้น (ตอนเล่น)")
table(["ชั้น", "จัดการ", "ความเร็ว"], [
    ["1. Cache (exact)", "ชื่อ/ปุ่ม/menu/affix", "<5ms"],
    ["2. Normalized", "ตัด space/พิมพ์ใหญ่เล็ก", "<5ms"],
    ["3. Template", "\"Defeat the X\", \"+12 Str\"", "<5ms"],
    ["4. Fuzzy", "กู้ OCR ที่อ่านพลาด", "~30ms"],
    ["5. Live (4B LLM)", "description ใหม่ → แปลสด + จำ", "~1วิ"],
], widths=[2.2, 3.6, 1.2])

heading("📂 โครงสร้างไฟล์สำคัญ")
table(["ไฟล์", "ทำอะไร"], [
    ["PLAY.bat", "กดเดียวเปิดทุกอย่างเพื่อเล่น"],
    ["matcher.py", "5 ชั้นค้นหา (exact/norm/template/fuzzy/miss)"],
    ["template.py", "จับ dynamic text ({MONSTER}, +12)"],
    ["server.py", "FastAPI ปลอม Ollama + live fallback"],
    ["cache.py", "SQLite พจนานุกรม (resume-able)"],
    ["glossary.yaml", "68 คำศัพท์เฉพาะ D4 (Lilith→ลิลิธ)"],
    ["cache.db", "96k+ คำแปล (หัวใจระบบ)"],
    ["translate_batch.py", "แปลทั้งหมด (parallel, resume)"],
    ["upgrade_live.py", "ยกคุณภาพคำที่ live เจอ (4B→8B)"],
], widths=[2.4, 4.4])

heading("📚 คำศัพท์เทคนิค (สำหรับมือใหม่)")
table(["คำ", "ความหมายง่ายๆ"], [
    ["CASC", "ระบบ pack ไฟล์ของ Blizzard (เหมือน zip ยักษ์)"],
    ["OCR", "อ่านตัวอักษรจากภาพ"],
    ["overlay", "หน้าต่างโปร่งใสวาดทับจอ"],
    ["LLM", "AI ภาษา เช่น Typhoon"],
    ["VRAM", "RAM ของการ์ดจอ — เกม+AI แย่งกันใช้"],
    ["cache", "ที่เก็บผลลัพธ์ไว้ใช้ซ้ำ (พจนานุกรม)"],
    ["fuzzy match", "จับคู่แบบใกล้เคียงพอ (กู้ OCR พลาด)"],
    ["template", "แม่แบบที่มีช่องว่าง (\"Defeat the ___\")"],
    ["memoize", "จำผลลัพธ์ ไม่คำนวณซ้ำ"],
    ["dual-stack", "socket รับทั้ง IPv4 และ IPv6"],
], widths=[1.8, 5.0])

heading("🌍 ถ้าจะทำเกมอื่น / ภาษาอื่น")
para("ใช้ซ้ำได้ ~70%: bridge, matcher, template, cache, batch translator, overlay (RST)")
para("ต้องทำใหม่ต่อเกม: extractor (ขึ้นกับ engine), glossary, dictionary")
para("เปลี่ยนภาษา: แค่เปลี่ยน target language ตอน batch + ฟอนต์ overlay")
para("เกม Unity/Unreal (single-player) มักแก้ไฟล์ตรงได้เลย (ไม่มี anti-cheat) — ง่ายกว่า D4", italic=True)

heading("✅ สถานะสุดท้าย")
for b in [
    "96,000+ คำแปลไทยใน cache",
    "เล่น D4 เห็นไทย real-time (UI/menu/skill/quest/item)",
    "live fallback เติมส่วนที่ขาด (~1วิ แล้วจำ)",
    "ปลอดภัย ไม่แตะไฟล์เกม ไม่เสี่ยง ban",
    "46 unit tests + CI เขียว",
    "repo: github.com/MeowV1101/babelmeow",
]:
    bullet(b)
para("")
para("จาก idea เล่นๆ → ระบบที่ใช้งานได้จริง 🐱", bold=True, size=13,
     color=RGBColor(0x1F,0x4E,0x79), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(str(OUT))
print(f"Saved: {OUT}  ({OUT.stat().st_size:,} bytes)")
